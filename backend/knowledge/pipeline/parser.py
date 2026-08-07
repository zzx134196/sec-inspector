"""文档解析器 — 支持PDF（国标法规）和Word文档"""
import os
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

from loguru import logger


def parse_document(file_path: str) -> str:
    """
    解析文档为纯文本
    :param file_path: 文档路径
    :return: 解析后的纯文本
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    file_ext = Path(file_path).suffix.lower()

    if file_ext == ".pdf":
        text = _parse_pdf(file_path)
    elif file_ext == ".docx":
        text = _parse_docx(file_path)
    elif file_ext in (".doc", ".wps"):
        text = _parse_doc(file_path)
    elif file_ext == ".txt":
        text = _parse_txt(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {file_ext}")

    text = _clean_parsed_text(text)
    return text


def _clean_parsed_text(text: str) -> str:
    """清理解析后的文本，移除不可打印字符"""
    if not text:
        return text
    cleaned = ''.join(
        ch for ch in text
        if ch.isprintable() or ch in ('\n', '\r', '\t')
    )
    if len(cleaned) < len(text) * 0.5:
        logger.warning(f"解析文本含大量不可打印字符({len(text) - len(cleaned)}个已移除)，可能为扫描件PDF")
    return cleaned


def _parse_pdf(file_path: str) -> str:
    """解析PDF文件 — 优先PyMuPDF，回退PyPDF2"""
    # PyMuPDF（fitz）解析质量更高，适合国标PDF
    try:
        import fitz
        text_parts = []
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text.strip())
        if text_parts:
            result = "\n\n".join(text_parts)
            logger.info(f"PyMuPDF解析成功: {file_path}, {len(result)}字符, {len(text_parts)}页")
            return result
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"PyMuPDF解析失败，尝试PyPDF2: {e}")

    # 回退到PyPDF2
    try:
        import PyPDF2
        text_parts = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        result = "\n\n".join(text_parts)
        logger.info(f"PyPDF2解析成功: {file_path}, {len(result)}字符")
        return result
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"PyPDF2解析也失败: {e}")

    raise Exception(f"PDF解析失败: 需要安装 PyMuPDF(fitz) 或 PyPDF2")


def _parse_docx(file_path: str) -> str:
    """解析.docx文件"""
    try:
        from docx import Document
        doc = Document(file_path)

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)
    except ImportError:
        raise ImportError("需要安装python-docx: pip install python-docx")


def _parse_doc(file_path: str) -> str:
    """解析旧版.doc文件 — 使用macOS textutil或LibreOffice转换"""
    # macOS textutil转txt
    try:
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            tmp_path = tmp.name

        result = subprocess.run(
            ["textutil", "-convert", "txt", "-output", tmp_path, file_path],
            capture_output=True, text=True, timeout=30,
        )

        if result.returncode == 0 and os.path.exists(tmp_path):
            with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read().strip()
            os.unlink(tmp_path)
            if text:
                return text
        else:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    except FileNotFoundError:
        pass
    except Exception:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    # LibreOffice转换
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "txt:Text", "--outdir", tmpdir, file_path],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0:
                txt_files = list(Path(tmpdir).glob("*.txt"))
                if txt_files:
                    with open(txt_files[0], "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read().strip()
                    if text:
                        return text
    except FileNotFoundError:
        pass

    raise Exception(f"旧版.doc解析失败: 无法找到可用的转换工具")


def _parse_txt(file_path: str) -> str:
    """解析纯文本文件"""
    for encoding in ("utf-8", "gbk", "gb2312", "gb18030"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise Exception(f"TXT文件解析失败: 无法识别编码")
