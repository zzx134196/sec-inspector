"""文档导出API — Word/PDF/Excel生成与下载"""
import io
import os
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from loguru import logger

from app.models.database import get_db
from app.models.user import User
from app.core.auth import get_current_user
from app.config import settings

router = APIRouter(prefix="/api/export", tags=["文档导出"])


class ExportWordRequest(BaseModel):
    title: str
    content: str


class ExportExcelRequest(BaseModel):
    columns: list
    rows: list
    sheet_name: str = "查询结果"


@router.post("/word")
async def export_word(
    req: ExportWordRequest,
    current_user: User = Depends(get_current_user),
):
    """导出Word文档"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 标题
        title_para = doc.add_heading(req.title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 正文
        for paragraph in req.content.split('\n'):
            stripped = paragraph.strip()
            if not stripped:
                doc.add_paragraph('')
                continue

            if stripped.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、')):
                doc.add_heading(stripped, level=2)
            elif stripped.startswith(('（一）', '（二）', '（三）', '（四）', '（五）')):
                doc.add_heading(stripped, level=3)
            else:
                p = doc.add_paragraph(stripped)
                p.style.font.size = Pt(12)

        doc.add_paragraph('')
        footer = doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
        footer.style.font.size = Pt(9)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        from urllib.parse import quote
        filename = f"{req.title}_{datetime.now().strftime('%Y%m%d')}.docx"
        encoded_filename = quote(filename)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
        )
    except Exception as e:
        logger.error(f"Word导出失败: {e}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.post("/pdf")
async def export_pdf(
    req: ExportWordRequest,
    current_user: User = Depends(get_current_user),
):
    """导出PDF文档"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_CENTER

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=25*mm, bottomMargin=25*mm)

        styles = getSampleStyleSheet()
        font_name = 'Helvetica'
        font_candidates = [
            '/System/Library/Fonts/STHeiti Light.ttc',
            '/System/Library/Fonts/PingFang.ttc',
            '/Library/Fonts/Arial Unicode.ttf',
            '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
            '/usr/share/fonts/wqy-zenhei/wqy-zenhei.ttc',
            '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
            'C:/Windows/Fonts/simhei.ttf',
        ]
        for font_path in font_candidates:
            try:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('CJKFont', font_path))
                    font_name = 'CJKFont'
                    break
            except Exception:
                continue

        title_style = ParagraphStyle('ChTitle', parent=styles['Title'], fontName=font_name, fontSize=18, alignment=TA_CENTER)
        body_style = ParagraphStyle('ChBody', parent=styles['Normal'], fontName=font_name, fontSize=12, leading=20)
        heading_style = ParagraphStyle('ChHeading', parent=styles['Heading2'], fontName=font_name, fontSize=14)

        story = []
        story.append(Paragraph(req.title, title_style))
        story.append(Spacer(1, 12))

        for line in req.content.split('\n'):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 6))
                continue
            if stripped.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、')):
                story.append(Paragraph(stripped, heading_style))
            else:
                story.append(Paragraph(stripped, body_style))

        doc.build(story)
        buffer.seek(0)

        from urllib.parse import quote
        filename = f"{req.title}_{datetime.now().strftime('%Y%m%d')}.pdf"
        encoded_filename = quote(filename)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
        )
    except Exception as e:
        logger.error(f"PDF导出失败: {e}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.post("/excel")
async def export_excel(
    req: ExportExcelRequest,
    current_user: User = Depends(get_current_user),
):
    """导出Excel文件"""
    try:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = req.sheet_name

        # 写入表头
        ws.append(req.columns)

        # 写入数据
        for row in req.rows:
            if isinstance(row, dict):
                ws.append([row.get(col, "") for col in req.columns])
            elif isinstance(row, (list, tuple)):
                ws.append(list(row))

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        from urllib.parse import quote
        filename = f"{req.sheet_name}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        encoded_filename = quote(filename)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
        )
    except Exception as e:
        logger.error(f"Excel导出失败: {e}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")
