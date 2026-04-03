"""生成等保测评报告测试样例docx文件"""
from docx import Document

doc = Document()
doc.add_heading('信息系统等级保护测评报告', level=0)

# 一、基本信息
doc.add_heading('一、被测系统基本信息', level=1)
t = doc.add_table(rows=7, cols=2, style='Table Grid')
for i, (k, v) in enumerate([
    ('系统名称', '某市政务服务管理平台'),
    ('系统编号', 'GP-2026-0315'),
    ('安全保护等级', '第三级'),
    ('测评机构', '某信息安全测评中心'),
    ('测评日期', '2026年3月15日'),
    ('被测单位', '某市大数据管理局'),
    ('系统简介', '该系统为某市统一政务服务平台，承载政务审批、公共服务等核心业务，涉及大量公民个人信息和政务数据。'),
]):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v

# 二、测评范围
doc.add_heading('二、测评范围与对象', level=1)
doc.add_paragraph('本次测评涵盖以下测评对象：')
for obj in [
    '应用服务器：CentOS 7.9（2台），IP: 192.168.1.10, 192.168.1.11',
    '数据库服务器：MySQL 8.0（1台），IP: 192.168.1.20',
    'Web应用服务器：Nginx 1.22 + Tomcat 9.0',
    '网络设备：华为S5700交换机（2台）、防火墙USG6000（1台）',
    '安全设备：入侵检测系统、日志审计系统',
]:
    doc.add_paragraph(obj, style='List Bullet')

# 三、安全控制测评结果
doc.add_heading('三、安全控制测评结果', level=1)

# 3.1 身份鉴别
doc.add_heading('3.1 身份鉴别', level=2)
doc.add_paragraph('【测评对象】CentOS 7.9 Linux服务器')
doc.add_paragraph(
    '【测评项】a) 应对登录的用户进行身份标识和鉴别，身份标识具有唯一性，'
    '身份鉴别信息具有复杂度要求并定期更换。'
)
doc.add_paragraph('【符合情况】部分符合')
doc.add_paragraph(
    '【测评描述】\n'
    '1. 经核查，系统通过用户名和密码方式对登录用户进行身份鉴别。'
    '执行 cat /etc/shadow 查看，所有账户均设置了密码。\n'
    '2. 执行 cat /etc/login.defs 查看密码策略：PASS_MAX_DAYS=90, '
    'PASS_MIN_DAYS=0, PASS_MIN_LEN=8, PASS_WARN_AGE=7。密码复杂度基本满足要求。\n'
    '3. 但经核查发现存在共享账户问题：运维团队5人共用root账户登录服务器，'
    '未配置独立账户，身份标识不具有唯一性。\n'
    '4. 未启用双因素认证机制。'
)

doc.add_paragraph(
    '【测评项】b) 应具有登录失败处理功能，应配置并启用结束会话、'
    '限制非法登录次数和当登录连接超时自动退出等相关措施。'
)
doc.add_paragraph('【符合情况】不符合')
doc.add_paragraph(
    '【测评描述】\n'
    '1. 经核查 /etc/pam.d/system-auth 文件，未配置 pam_tally2 或 pam_faillock 模块，'
    '未限制登录失败次数。\n'
    '2. 执行 cat /etc/ssh/sshd_config 查看，ClientAliveInterval=0, '
    'ClientAliveCountMax=3，未启用登录连接超时自动退出功能。\n'
    '3. 由于业务连续性要求，运维人员反映不能限制登录失败次数，'
    '担心被锁定后影响紧急运维操作。'
)

# 3.2 访问控制
doc.add_heading('3.2 访问控制', level=2)
doc.add_paragraph('【测评对象】CentOS 7.9 Linux服务器')
doc.add_paragraph('【测评项】a) 应对登录的用户分配账户和权限。')
doc.add_paragraph('【符合情况】部分符合')
doc.add_paragraph(
    '【测评描述】\n'
    '1. 系统已为不同角色创建了相应账户，包括管理员账户和普通用户账户。\n'
    '2. 但核查发现，普通运维人员账户被赋予了sudoers权限'
    '（ALL=(ALL) NOPASSWD:ALL），可以无密码执行任意root命令，违反最小权限原则。\n'
    '3. 存在多个已离职员工的账户未及时清理，账户zhangwei、liuqiang仍处于活跃状态。'
)

# 3.3 安全审计
doc.add_heading('3.3 安全审计', level=2)
doc.add_paragraph('【测评对象】CentOS 7.9 Linux服务器')
doc.add_paragraph(
    '【测评项】a) 应启用安全审计功能，审计覆盖到每个用户，'
    '对重要的用户行为和重要安全事件进行审计。'
)
doc.add_paragraph('【符合情况】符合')
doc.add_paragraph(
    '【测评描述】\n'
    '1. 系统已安装并启用rsyslog服务，执行 systemctl status rsyslog 显示active (running)。\n'
    '2. 已配置auditd审计服务，审计规则覆盖关键文件的读写操作和用户登录行为。\n'
    '3. 审计日志保存在本地/var/log/目录下，同时通过syslog转发至日志审计系统集中存储。\n'
    '4. 日志保存期限已配置为180天，满足等保三级要求。'
)

# 3.4 数据库安全
doc.add_heading('3.4 数据库安全', level=2)
doc.add_paragraph('【测评对象】MySQL 8.0数据库')
doc.add_paragraph('【测评项】a) 应对登录的用户进行身份标识和鉴别。')
doc.add_paragraph('【符合情况】部分符合')
doc.add_paragraph(
    '【测评描述】\n'
    '1. MySQL数据库通过用户名和密码进行身份鉴别，已启用validate_password插件。\n'
    '2. 但数据库root账户允许从任意主机（%）登录，未限制登录来源IP。\n'
    '3. 应用程序连接数据库使用的是root账户，而非专用的应用账户，权限过大。\n'
    '4. 数据库密码为admin@123，复杂度不足，且超过6个月未更换。'
)

# 四、总体结论
doc.add_heading('四、总体测评结论', level=1)
doc.add_paragraph('本次测评发现信息系统在安全控制方面存在以下主要问题：')
for issue in [
    '【高风险】多人共用root账户，身份标识不唯一',
    '【高风险】未配置登录失败锁定策略',
    '【高风险】数据库root账户允许任意IP访问，密码强度不足',
    '【中风险】运维账户权限过大，违反最小权限原则',
    '【中风险】存在离职员工账户未及时清理',
    '【低风险】未启用SSH连接超时退出',
]:
    doc.add_paragraph(issue, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph(
    '综合以上测评结果，该信息系统目前不满足等级保护第三级安全要求。'
    '建议被测单位按照上述问题进行整改后，申请复测。'
)

import os
path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '等保测评报告_测试样例.docx')
doc.save(path)
print(f'文件已生成: {path}')
